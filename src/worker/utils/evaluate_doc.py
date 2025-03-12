import re
import yaml
from dataclasses import dataclass
from docx import Document
from typing import List, Optional
from worker.utils.genext import GenextAPI, LlmApiModel
import logging


# Configure logger
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)



@dataclass
class DocNode:
    text: str
    style: str
    level: int
    children: List['DocNode']
    parent: Optional['DocNode'] = None


class DocParser:
    def __init__(self):
        self.heading_pattern = re.compile(r'Heading (\d+)')
        self.root = DocNode("", "root", 0, [])

    def parse_document(self, file_path: str) -> DocNode:
        doc = Document(file_path)
        current_node = self.root
        current_level = 0
        counter = 0
        paragraph_nr = 0
        table_nr = 0
        # loop through all the items
        for element in doc.element.body:
            counter = counter + 1

            if element.tag.endswith('p'):
                # print(counter, 'P', paragraph_nr)

                paragraph = doc.paragraphs[paragraph_nr]
                style = paragraph.style.name
                heading_match = self.heading_pattern.match(style)

                if heading_match:
                    level = int(heading_match.group(1))
                    node = DocNode(paragraph.text, style, level, [])

                    # Move up the tree if needed
                    while current_level >= level and current_node.parent:
                        current_node = current_node.parent
                        current_level -= 1

                    # Add new node as child of current node
                    node.parent = current_node
                    current_node.children.append(node)
                    current_node = node
                    current_level = level
                else:
                    # Regular paragraph - add as child of current section
                    node = DocNode(paragraph.text, style, current_level + 1, [])
                    # node = DocNode(paragraph.text, 'Normal', current_level + 1, [])
                    node.parent = current_node
                    current_node.children.append(node)
                paragraph_nr = paragraph_nr + 1

            elif element.tag.endswith('tbl'):
                # print(counter, 'T', table_nr)

                row_nr = 0

                table = doc.tables[table_nr]
                table_data = []
                for row in table.rows:
                    row_nr = row_nr + 1
                    table_name = 'Table_' + str(table_nr) + '_Row_' + str(row_nr)
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                    node = DocNode(str(row_data), table_name, current_level + 1, [])
                    node.parent = current_node
                    current_node.children.append(node)

                table_nr = table_nr + 1

            # else:
            # print(counter, 'O')


        return self.root

    def traverse_depth_first(self, node: DocNode = None, level: int = 0):
        if node is None:
            node = self.root

        yield (node, level)
        for child in node.children:
            yield from self.traverse_depth_first(child, level + 1)

    def find_sections_by_style(self, style: str) -> List[DocNode]:
        return [node for node, _ in self.traverse_depth_first()
                if node.style == style]

    def find_text(self, search_text: str) -> List[DocNode]:
        return [node for node, _ in self.traverse_depth_first()
                if search_text.lower() in node.text.lower()]

    ####

    def find_text_with_subnodes(self, search_text: str) -> List[DocNode]:
        matching_nodes = []
        for node, _ in self.traverse_depth_first():
            if search_text.lower() in node.text.lower():
                matching_nodes.append(node)
                matching_nodes.extend(self.get_subnodes(node))
        return matching_nodes

    def get_subnodes(self, node: DocNode) -> List[DocNode]:
        subnodes = []
        for child in node.children:
            subnodes.append(child)
            subnodes.extend(self.get_subnodes(child))
        return subnodes


class ValidationRules:
    def __init__(self):
        self.ifile = 1

    def read_yaml_file(self, file_path):
        with open(file_path, 'r') as file:
            return yaml.safe_load(file)

    def create_list_of_rules(self, rule_data):
        # Loop through sections and validation rules
        lv_sections = []
        lv_rules = []
        lv_rules_ids = []

        for section in rule_data.get('sections', []):
            lv_section = section.get('section_text')

            for rule in section.get('validation_rules', []):
                lv_rule = 'Rule ID: ' + rule.get('rule_id') + '\n'
                lv_rule = lv_rule + 'Prompt: ' + rule.get('prompt') + '\n'
                lv_rule = lv_rule + 'Criteria: \n'
                lv_rule_id = rule.get('rule_id')

                for criterion in rule.get('criteria', []):
                    lv_rule = lv_rule + '    - ' + criterion + '\n'

                lv_sections.append(lv_section)
                lv_rules.append(lv_rule)
                lv_rules_ids.append(lv_rule_id)

        return lv_sections, lv_rules, lv_rules_ids


def extract_score(response):
    score = 0
    lv_response = response
    lv_response = lv_response.replace('Score: ', '')
    lv_response = lv_response.replace('score: ', '')
    lv_response = lv_response.replace('Rating: ', '')
    lv_response = lv_response.replace('rating: ', '')

    while True:
        try:
            # Extract the first number from the response string
            match = re.search(r'\d+', lv_response)
            if match:
                score = int(match.group(0))
                break  # Exit the loop if a number is found
            else:
                print('ERROR!!!: ' + lv_response)
                score = 0
        except ValueError:
            print('ERROR!!!: ' + lv_response)
        # score = 5  # or any default value you prefer

    return score


def do_evaluation(document_path, yaml_file_path):
    # read the document
    parser = DocParser()
    doc_tree = parser.parse_document(document_path)

    # Print document structure and put in a string
    document_text = ''
    for node, level in parser.traverse_depth_first():
        if node.text != '':
            document_text = document_text + ("  " * level + f"- {node.text}\n")
    # read yaml and prepare the prompt
    valRules = ValidationRules()
    rule_data = valRules.read_yaml_file(yaml_file_path)
    sections, rules, rules_id = valRules.create_list_of_rules(rule_data)

    # evaluate the document
    index = 0
    lt_score = []
    lt_answer = []

    for rule in rules:

        # get the content
        content = ('TASK: You evaluate documents based on evaluation criteria and provide a rating out of 10.' +
                   'You first provide a number out of 10 for the score then provide feedback based on this evaluation criteria:' +
                #    'Section Text: Provide the section_text as is in the output to help with referencing. Add the perfix section_text: ' +
                   'EVALUATION CRITERIA: ' +
                    rule +
                   'SCORING: The scoring is out of 10 and you always give a score. For example:  \'10 - The document seems complete.\'' +
                   'Or: \'5 - Your document still needs work, for example it is missing contact information. \'' +
                   'Or: \'0 - Your document is not good enough, please revisit. \'')

        # get the prompt
        section = sections[index]
        rule_id = rules_id[index]

        # print('-------------------------------------------------------------------')
        # print('Rule Id: ' + rule_id)

        index = index + 1

        # no section then use the whole document
        if section != '':
            # print(' - ' + section)
            results = parser.find_text_with_subnodes(section)
            document_text_sub = ''
            for result in results:
                document_text_sub = document_text_sub + result.text + '\n'
            prompt = 'Please evaluate this document: "' + document_text_sub + '"'
            # print('Sub text' + document_text_sub)
        else:
            # print(' - Complete Document')
            prompt = 'Please evaluate this document: ' + document_text
            # print(document_text)

        # call the API
        genext_api = GenextAPI(
            question=prompt,
            model_name=LlmApiModel.GPT_4o,
            temperature=0.2,
            max_completion_token_count=400,
            content=content
        )

        # print the feedback
        response = genext_api.run()
        answer = response['completion']
        # print(' - ' + answer + '\n')
        score = extract_score(answer)
        # print('Score: ' + str(score))

        lt_score.append(score)
        lt_answer.append(answer)

    return lt_score, lt_answer


def perform_evaluation(document_path, yaml_file_path):
    # document_path = '/home/qxz1viq/doc_eval_latest/evaluation_processor/data/example_doc.docx'
    # yaml_file_path = '/home/qxz1viq/doc_eval_latest/evaluation_processor/data/checklist.yaml'

    val_rules = ValidationRules()
    rule_data = val_rules.read_yaml_file(yaml_file_path)
    sections, rules, rules_id = val_rules.create_list_of_rules(rule_data)
    

    lt_score, lt_answer = do_evaluation(document_path, yaml_file_path)

    findings = []
    for section, score, answer in zip(sections, lt_score, lt_answer):
        findings.append({
            "section_name": section,
            "summary": f"Score: {score} for section {section}",
            "score": score,
            "details": answer
        })

    print(findings)
    return findings


# read document and evaluate
if __name__ == "__main__":

    # input
    # document_path = 'input_files/emergency.docx'
    # yaml_file_path = 'input_files/document_evaluation.yml'
    document_path = '/home/qxz1viq/doc_eval_latest/evaluation_processor/data/example_doc.docx'
    yaml_file_path = '/home/qxz1viq/doc_eval_latest/evaluation_processor/data/checklist.yaml'
    perform_evaluation(document_path, yaml_file_path)


